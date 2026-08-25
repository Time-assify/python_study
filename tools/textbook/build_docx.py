# -*- coding: utf-8 -*-
"""讲义docx渲染器：统一模板 → Word文档

模板（每天九节）:
  1 今日学习目标 / 2 为什么需要学习这个 / 3 核心知识讲解 / 4 图示·流程
  5 最小代码示例 / 6 工程实践 / 7 常见错误 / 8 今日总结 / 9 与今日任务关系
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

EAST_ASIA_FONT = "微软雅黑"
BODY_FONT = "宋体"
CODE_FONT = "Consolas"

_BOLD_RE = re.compile(r"(\*\*.+?\*\*)")


def _set_east_asia(style_or_run, font_name):
    rpr = style_or_run.font._element.get_or_add_rPr() \
        if hasattr(style_or_run, "font") else style_or_run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def _setup_styles(doc):
    # 标题样式：中文黑体感
    for name, size in (("Heading 1", 20), ("Heading 2", 16), ("Heading 3", 13)):
        st = doc.styles[name]
        st.font.name = EAST_ASIA_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x1F, 0x3B, 0x63)
        _set_east_asia(st, EAST_ASIA_FONT)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    _set_east_asia(normal, BODY_FONT)

    if "CodeBlock" not in [s.name for s in doc.styles]:
        code = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = CODE_FONT
        code.font.size = Pt(9.5)
        code.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        code.paragraph_format.space_after = Pt(0)
        code.paragraph_format.space_before = Pt(0)
        ppr = code._element.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "F2F2F2")
        ppr.append(shd)

    if "Caption" not in [s.name for s in doc.styles]:
        cap = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = BODY_FONT
        cap.font.size = Pt(9)
        cap.font.italic = True
        cap.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        _set_east_asia(cap, BODY_FONT)


def _add_runs(paragraph, text):
    """支持 **加粗** 标记"""
    for part in _BOLD_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
    return paragraph


def add_rich(doc, text, style=None, indent=False):
    p = doc.add_paragraph(style=style)
    _add_runs(p, text)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    return p


def add_bullet(doc, text):
    return add_rich(doc, text, style="List Bullet")


def add_code(doc, code, title=None, caption=None):
    cap_text = title or caption
    if cap_text:
        cap = doc.add_paragraph(style="Caption")
        cap.add_run(cap_text)
    for line in code.rstrip("\n").split("\n"):
        p = doc.add_paragraph(style="CodeBlock")
        p.add_run(line)
        p.paragraph_format.left_indent = Inches(0.15)
    doc.add_paragraph()  # 间隔


def add_image(doc, path, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(4.6))
    if caption:
        cap = doc.add_paragraph(style="Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run(caption)


def add_toc(doc):
    """插入可自动更新的目录域（Word中右键→更新域，或F9）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_runs(p, "**目  录**")
    note = doc.add_paragraph(style="Caption")
    note.add_run("（在Word中右键此处 → 更新域，即可生成页码；或用 F9）")
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._element.append(fld)
    run2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    run2._element.append(instr)
    run3 = p.add_run()
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "separate")
    run3._element.append(fld2)
    run4 = p.add_run("（打开后更新域显示目录）")
    run5 = p.add_run()
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "end")
    run5._element.append(fld3)
    doc.add_page_break()


def render_mistake_table(doc, mistakes):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, t in enumerate(("问题", "原因", "解决")):
        hdr[i].paragraphs[0].add_run(t).bold = True
    for m in mistakes:
        cells = table.add_row().cells
        for i, key in enumerate(("q", "reason", "fix")):
            cells[i].paragraphs[0].add_run(m[key])
    doc.add_paragraph()


def render_day(doc, content):
    doc.add_heading(f"Day {content['day']:02d}: {content['title']}", level=2)

    doc.add_heading("1. 今日学习目标", level=3)
    add_rich(doc, content["goal"])

    doc.add_heading("2. 为什么需要学习这个", level=3)
    for para in content["why"]:
        if para.startswith("- "):
            add_bullet(doc, para[2:])
        else:
            add_rich(doc, para)

    doc.add_heading("3. 核心知识讲解", level=3)
    for item in content["core"]:
        if isinstance(item, tuple):
            subtitle, paras = item
            p = add_rich(doc, f"**{subtitle}**")
            p.paragraph_format.space_before = Pt(6)
            for para in paras:
                if para.startswith("- "):
                    add_bullet(doc, para[2:])
                else:
                    add_rich(doc, para, indent=True)
        elif item.startswith("- "):
            add_bullet(doc, item[2:])
        else:
            add_rich(doc, item)

    doc.add_heading("4. 图示 · 流程", level=3)
    for fig in content["diagrams"]:
        kind = fig[0]
        if kind == "ascii":
            add_code(doc, fig[1], caption=fig[2] if len(fig) > 2 else None)
        elif kind == "image":
            path = fig[1]() if callable(fig[1]) else fig[1]
            add_image(doc, path, caption=fig[2] if len(fig) > 2 else None)
        elif kind == "para":
            add_rich(doc, fig[1])

    doc.add_heading("5. 最小代码示例", level=3)
    for snippet in content["code"]:
        title, code, explains = snippet
        add_code(doc, code, title=title)
        for e in explains:
            add_rich(doc, e, indent=True)

    doc.add_heading("6. 工程实践", level=3)
    for para in content["practice"]:
        if para.startswith("- "):
            add_bullet(doc, para[2:])
        else:
            add_rich(doc, para)

    doc.add_heading("7. 常见错误", level=3)
    render_mistake_table(doc, content["mistakes"])

    doc.add_heading("8. 今日总结", level=3)
    add_rich(doc, "**今天掌握**")
    for item in content["summary"]["learned"]:
        add_bullet(doc, item)
    add_rich(doc, "**必须会**")
    for item in content["summary"]["must"]:
        add_bullet(doc, item)

    doc.add_heading("9. 与今日任务关系", level=3)
    add_rich(doc, content["task_link"])


def render_part(doc, part):
    doc.add_heading(part["title"], level=1)
    add_rich(doc, part["intro"])
    for day in part["days"]:
        render_day(doc, day)
        doc.add_page_break()


def build_document(output_path, parts):
    doc = Document()
    _setup_styles(doc)

    # 封面
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(180)
    _add_runs(p, "**Python Study**")
    for r in p.runs:
        r.font.size = Pt(40)
        r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x63)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_runs(p2, "**AI工程师 40天课程讲义**")
    for r in p2.runs:
        r.font.size = Pt(24)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_runs(p3, "从零基础到 PyTorch · 计算机视觉 · LLM应用")
    for r in p3.runs:
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_page_break()

    add_toc(doc)

    for part in parts:
        render_part(doc, part)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
